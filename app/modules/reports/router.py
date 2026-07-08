import logging

from fastapi import APIRouter, Depends, HTTPException, Request
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from app.core.dependencies import get_current_user
from app.db.mysql import get_db
from app.models import Report, User
from app.schemas import ReportCreateIn
from app.services.audit_service import audit_action
from app.services.generation_service import GenerationService
from app.utils.json_utils import loads

logger = logging.getLogger(__name__)
router = APIRouter(prefix='/reports', tags=['研读报告'])


def _serialize_report_summary(obj: Report) -> dict:
    """创建接口只返回摘要，避免超大 JSON 导致前端网络错误。"""
    return {
        'id': obj.id,
        'paper_id': obj.paper_id,
        'title': obj.title,
        'created_at': obj.created_at,
    }


@router.post('')
async def create_report(
    data: ReportCreateIn,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    try:
        obj = await GenerationService().create_report(db, user.id, data.paper_id, data.modules, data.title)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    # create_report 内部已 commit；此处仅补审计/计数，失败不应让前端误判为“报告未生成”
    try:
        await db.refresh(user)
        user.report_generate_count += 1
        await audit_action(
            db,
            user_id=user.id,
            module='reports',
            operation_type='generate',
            content={
                'report_id': obj.id,
                'paper_id': obj.paper_id,
                'title': obj.title,
                'modules': data.modules,
            },
            request=request,
        )
        await db.commit()
    except Exception:
        await db.rollback()
        logger.exception('Report %s saved but audit/quota update failed', obj.id)

    return _serialize_report_summary(obj)

@router.get('')
async def list_reports(db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    rows = (await db.execute(select(Report).where(Report.user_id == user.id).order_by(Report.created_at.desc()))).scalars().all()
    return [{'id': x.id, 'paper_id': x.paper_id, 'title': x.title, 'content': loads(x.content), 'created_at': x.created_at} for x in rows]

@router.delete('/{report_id}')
async def delete_report(
    report_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    report = await db.get(Report, report_id)
    if not report or report.user_id != user.id:
        raise HTTPException(404, '报告不存在')
    await audit_action(
        db,
        user_id=user.id,
        module='reports',
        operation_type='delete',
        content={'report_id': report.id, 'paper_id': report.paper_id, 'title': report.title},
        request=request,
        risk=1,
    )
    await db.delete(report)
    await db.commit()
    return {'message': 'deleted'}

from fastapi.responses import Response
from io import BytesIO
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from urllib.parse import quote


def _ascii_filename(title: str | None, ext: str) -> str:
    base = re.sub(r'[^A-Za-z0-9._-]+', '_', title or 'report').strip('._')
    if not base:
        base = 'report'
    return f'{base[:80]}.{ext}'


def _download_headers(title: str | None, ext: str) -> dict[str, str]:
    safe_title = (title or 'report').replace('/', '_').replace('\\', '_')
    utf8_name = f'{safe_title}.{ext}'
    ascii_name = _ascii_filename(safe_title, ext)
    return {
        'Content-Disposition': (
            f"attachment; filename={ascii_name}; "
            f"filename*=UTF-8''{quote(utf8_name)}"
        )
    }

def _plain_text(text: object) -> str:
    value = str(text or '')
    value = re.sub(r'\*\*(.*?)\*\*', r'\1', value)
    value = re.sub(r'\[(.*?)\]\((.*?)\)', r'\1 (\2)', value)
    value = value.replace('`', '')
    return value.strip()


def _display_export_value(value: object) -> str:
    text = _plain_text(value)
    if not text or text == '-':
        return '-'
    if text.startswith('匹配关键词'):
        return ''
    if '当前解析未提取' in text or '原文片段中未充分体现' in text or text.startswith('暂未'):
        return '-'
    return text


def _report_visual_summary_markdown(content: dict) -> str:
    visual = content.get('visual_summary') or {}
    if not isinstance(visual, dict):
        return ''

    lines: list[str] = ['## 速览', '']
    flow = visual.get('method_flow') or []
    if flow:
        lines.append('### 方法流程')
        lines.append('')
        for idx, item in enumerate(flow, start=1):
            title = item.get('title') or f'步骤 {idx}'
            body = _display_export_value(item.get('content'))
            lines.append(f'{idx}. **{title}**：{body}')
        lines.append('')

    table = visual.get('key_data_table') or []
    if table:
        lines.append('### 关键数据')
        lines.append('')
        lines.append('| 字段 | 内容 |')
        lines.append('| --- | --- |')
        for row in table:
            value = _display_export_value(row.get('value')).replace('|', '/')
            lines.append(f'| {row.get("name") or ""} | {value} |')
        lines.append('')

    cards = visual.get('metric_cards') or []
    lines.append('### 核心指标')
    lines.append('')
    if cards:
        for card in cards:
            extra = []
            if card.get('note'):
                extra.append(_display_export_value(card['note']))
            suffix = f'（{"；".join(extra)}）' if extra and extra[0] != '-' else ''
            value = _display_export_value(card.get('value'))
            lines.append(f'- **{card.get("label") or "指标"}**：{value}{suffix}')
    else:
        lines.append('-')
    return '\n'.join(lines).strip()


def _reference_links_markdown(content: dict) -> str:
    refs = content.get('reference_links') or []
    if not refs:
        return ''
    lines = ['## 延伸阅读', '']
    for idx, item in enumerate(refs, start=1):
        label = item.get('raw') or item.get('title') or f'参考文献 {idx}'
        url = item.get('url')
        if url:
            lines.append(f'{idx}. [{label}]({url})')
        else:
            lines.append(f'{idx}. {label}')
        reason = _display_export_value(item.get('reason'))
        if reason:
            lines.append(f'   - 说明：{reason}')
    return '\n'.join(lines).strip()


def _strip_report_h1(markdown: str) -> str:
    return re.sub(r'^\s*#\s+.+?\n+', '', markdown.strip(), count=1)


def _strip_reference_sections(markdown: str) -> str:
    text = re.sub(
        r'\n{2,}##\s*(?:文献溯源链接|可点击文献溯源)\s*\n[\s\S]*$',
        '',
        markdown.strip(),
    )
    text = re.sub(
        r'\n{2,}##\s*文献溯源\s*\n\s*可点击参考文献已在结构化文献溯源模块中展示；导出文件会保留完整链接列表。\s*$',
        '',
        text,
    )
    text = re.sub(
        r'\n?###\s*原文参考文献溯源\s*\n[\s\S]*?(?=\n###\s*基础知识与拓展检索式|\n##\s|\Z)',
        '\n',
        text,
    )
    return text.strip()


def _build_export_markdown(content: dict, title: str | None = None) -> str:
    markdown = content.get('markdown') or str(content)
    markdown = _strip_reference_sections(markdown)
    markdown = _strip_report_h1(markdown)
    visual = _report_visual_summary_markdown(content)
    parts = [f'# {title or "研读报告"}']
    if visual:
        parts.append(visual)
    if markdown:
        parts.append(f'## 正文研读报告\n\n{markdown}')
    refs = _reference_links_markdown(content)
    if refs:
        parts.append(refs)
    return '\n\n'.join(parts)


def _body_markdown(content: dict) -> str:
    markdown = content.get('markdown') or str(content)
    markdown = _strip_reference_sections(markdown)
    return _strip_report_h1(markdown)


def _set_docx_cell_text(cell, text: object) -> None:
    cell.text = _plain_text(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def _set_docx_default_font(doc: Document) -> None:
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def _set_table_widths(table, widths: list[int]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths[:len(row.cells)]):
            tc = row.cells[idx]._tc
            tc_pr = tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in('w:tcW')
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(width))
            tc_w.set(qn('w:type'), 'dxa')


def _add_key_data_table_to_docx(doc: Document, rows: list[dict]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    for idx, header in enumerate(['字段', '内容']):
        _set_docx_cell_text(table.rows[0].cells[idx], header)
    for row in rows:
        cells = table.add_row().cells
        _set_docx_cell_text(cells[0], row.get('name') or '')
        _set_docx_cell_text(cells[1], _display_export_value(row.get('value')))
    _set_table_widths(table, [2200, 8200])


def _add_visual_summary_to_docx(doc: Document, visual: dict) -> None:
    if not isinstance(visual, dict) or not visual:
        return
    doc.add_heading('速览', level=2)
    flow = visual.get('method_flow') or []
    if flow:
        doc.add_heading('方法流程', level=3)
        for idx, item in enumerate(flow, start=1):
            title = item.get('title') or f'步骤 {idx}'
            body = _display_export_value(item.get('content'))
            doc.add_paragraph(f'{title}：{body}', style='List Number')
    table = visual.get('key_data_table') or []
    if table:
        doc.add_heading('关键数据', level=3)
        _add_key_data_table_to_docx(doc, table)
    doc.add_heading('核心指标', level=3)
    cards = visual.get('metric_cards') or []
    if cards:
        for card in cards:
            extra = []
            if card.get('note'):
                note = _display_export_value(card['note'])
                if note != '-':
                    extra.append(note)
            suffix = f'（{"；".join(extra)}）' if extra else ''
            value = _display_export_value(card.get('value'))
            doc.add_paragraph(f'{card.get("label") or "指标"}：{value}{suffix}', style='List Bullet')
    else:
        doc.add_paragraph('-', style='List Bullet')


def _add_reference_links_to_docx(doc: Document, refs: list[dict]) -> None:
    if not refs:
        return
    doc.add_heading('延伸阅读', level=2)
    for item in refs:
        title = item.get('raw') or item.get('title') or '未命名参考文献'
        url = item.get('url')
        text = f'{title}：{url}' if url else title
        doc.add_paragraph(_plain_text(text), style='List Number')


def _add_report_docx(doc: Document, content: dict, title: str) -> None:
    _set_docx_default_font(doc)
    doc.add_heading(title, level=1)
    _add_visual_summary_to_docx(doc, content.get('visual_summary') or {})
    body = _body_markdown(content)
    if body:
        doc.add_heading('正文研读报告', level=2)
        _add_markdown_to_docx(doc, body)
    _add_reference_links_to_docx(doc, content.get('reference_links') or [])


def _markdown_table_rows(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip('|').split('|')]
        rows.append(cells)
    return rows


def _add_markdown_to_docx(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()
        if not stripped:
            idx += 1
            continue

        if stripped.startswith('|') and idx + 1 < len(lines) and re.match(r'^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$', lines[idx + 1]):
            table_lines = [stripped]
            idx += 2
            while idx < len(lines) and lines[idx].strip().startswith('|'):
                table_lines.append(lines[idx].strip())
                idx += 1
            rows = _markdown_table_rows(table_lines)
            if rows:
                table = doc.add_table(rows=1, cols=len(rows[0]))
                table.style = 'Table Grid'
                for col, cell in enumerate(rows[0]):
                    table.rows[0].cells[col].text = _plain_text(cell)
                for row in rows[1:]:
                    cells = table.add_row().cells
                    for col, cell in enumerate(row[:len(cells)]):
                        cells[col].text = _plain_text(cell)
            continue

        if stripped.startswith('# '):
            doc.add_heading(_plain_text(stripped[2:]), level=1)
        elif stripped.startswith('## '):
            doc.add_heading(_plain_text(stripped[3:]), level=2)
        elif stripped.startswith('### '):
            doc.add_heading(_plain_text(stripped[4:]), level=3)
        elif re.match(r'^\d+\.\s+', stripped):
            doc.add_paragraph(_plain_text(re.sub(r'^\d+\.\s+', '', stripped)), style='List Number')
        elif stripped.startswith(('- ', '* ')):
            doc.add_paragraph(_plain_text(stripped[2:]), style='List Bullet')
        else:
            doc.add_paragraph(_plain_text(stripped))
        idx += 1


def _wrap_text(text: str, max_chars: int = 56) -> list[str]:
    value = _plain_text(text) or ' '
    return [value[i:i + max_chars] for i in range(0, len(value), max_chars)] or [' ']


def _draw_pdf_line(c: canvas.Canvas, text: str, x: int, y: float, font_name: str, size: int, width_chars: int = 56) -> float:
    c.setFont(font_name, size)
    for line in _wrap_text(text, width_chars):
        if y < 52:
            c.showPage()
            c.setFont(font_name, size)
            y = A4[1] - 50
        c.drawString(x, y, line)
        y -= size + 6
    return y


def _draw_markdown_pdf(c: canvas.Canvas, markdown: str, font_name: str) -> None:
    y = A4[1] - 50
    lines = markdown.splitlines()
    idx = 0
    while idx < len(lines):
        stripped = lines[idx].strip()
        if not stripped:
            y -= 6
            idx += 1
            continue

        if stripped.startswith('|') and idx + 1 < len(lines) and re.match(r'^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$', lines[idx + 1]):
            table_lines = [stripped]
            idx += 2
            while idx < len(lines) and lines[idx].strip().startswith('|'):
                table_lines.append(lines[idx].strip())
                idx += 1
            for row in _markdown_table_rows(table_lines):
                y = _draw_pdf_line(c, ' | '.join(_plain_text(cell) for cell in row), 50, y, font_name, 9, 62)
            y -= 6
            continue

        if stripped.startswith('# '):
            y = _draw_pdf_line(c, _plain_text(stripped[2:]), 45, y, font_name, 15, 38)
            y -= 6
        elif stripped.startswith('## '):
            y = _draw_pdf_line(c, _plain_text(stripped[3:]), 45, y, font_name, 13, 44)
            y -= 4
        elif stripped.startswith('### '):
            y = _draw_pdf_line(c, _plain_text(stripped[4:]), 45, y, font_name, 11, 48)
        else:
            y = _draw_pdf_line(c, stripped, 50, y, font_name, 10, 54)
        idx += 1


def _pdf_styles(font_name: str) -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        'title': ParagraphStyle('ReportTitle', parent=base['Title'], fontName=font_name, fontSize=18, leading=24, spaceAfter=12),
        'h2': ParagraphStyle('ReportH2', parent=base['Heading2'], fontName=font_name, fontSize=14, leading=18, spaceBefore=10, spaceAfter=6),
        'h3': ParagraphStyle('ReportH3', parent=base['Heading3'], fontName=font_name, fontSize=11.5, leading=15, spaceBefore=8, spaceAfter=4),
        'body': ParagraphStyle('ReportBody', parent=base['BodyText'], fontName=font_name, fontSize=9.5, leading=15, spaceAfter=5),
        'small': ParagraphStyle('ReportSmall', parent=base['BodyText'], fontName=font_name, fontSize=8.5, leading=13, spaceAfter=3),
    }


def _para(text: object, style: ParagraphStyle) -> Paragraph:
    value = _plain_text(text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    return Paragraph(value or ' ', style)


def _build_pdf_story(content: dict, title: str, font_name: str) -> list:
    styles = _pdf_styles(font_name)
    story = [_para(title, styles['title'])]
    visual = content.get('visual_summary') or {}
    if isinstance(visual, dict) and visual:
        story.extend([_para('速览', styles['h2']), _para('方法流程', styles['h3'])])
        flow_items = []
        for idx, item in enumerate(visual.get('method_flow') or [], start=1):
            title_text = item.get('title') or f'步骤 {idx}'
            body = _display_export_value(item.get('content'))
            flow_items.append(ListItem(_para(f'{title_text}：{body}', styles['body'])))
        if flow_items:
            story.append(ListFlowable(flow_items, bulletType='1', leftIndent=16))
        else:
            story.append(_para('-', styles['body']))

        table_rows = visual.get('key_data_table') or []
        if table_rows:
            story.append(_para('关键数据', styles['h3']))
            data = [[_para('字段', styles['small']), _para('内容', styles['small'])]]
            for row in table_rows:
                data.append([
                    _para(row.get('name'), styles['small']),
                    _para(_display_export_value(row.get('value')), styles['small']),
                ])
            table = Table(data, colWidths=[40 * mm, 138 * mm], repeatRows=1)
            table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), font_name),
                ('FONTSIZE', (0, 0), (-1, -1), 8.5),
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#eef3f8')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#1e293b')),
                ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#d8e2ec')),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 5),
                ('RIGHTPADDING', (0, 0), (-1, -1), 5),
            ]))
            story.extend([table, Spacer(1, 6)])

        story.append(_para('核心指标', styles['h3']))
        cards = visual.get('metric_cards') or []
        metric_items = []
        if cards:
            for card in cards:
                note = _display_export_value(card.get('note'))
                extra = f'（{note}）' if note and note != '-' else ''
                value = _display_export_value(card.get('value'))
                metric_items.append(ListItem(_para(f'{card.get("label") or "指标"}：{value}{extra}', styles['body'])))
        else:
            metric_items.append(ListItem(_para('-', styles['body'])))
        story.append(ListFlowable(metric_items, bulletType='bullet', leftIndent=16))

    body = _body_markdown(content)
    if body:
        story.append(_para('正文研读报告', styles['h2']))
        for raw in body.splitlines():
            line = raw.strip()
            if not line:
                continue
            if line.startswith('## '):
                story.append(_para(line[3:], styles['h2']))
            elif line.startswith('### '):
                story.append(_para(line[4:], styles['h3']))
            elif line.startswith('# '):
                story.append(_para(line[2:], styles['h2']))
            elif line.startswith(('- ', '* ')):
                story.append(_para(f'• {line[2:]}', styles['body']))
            else:
                story.append(_para(line, styles['body']))

    refs = content.get('reference_links') or []
    if refs:
        story.append(_para('延伸阅读', styles['h2']))
        for idx, item in enumerate(refs, start=1):
            label = item.get('raw') or item.get('title') or f'参考文献 {idx}'
            url = item.get('url')
            story.append(_para(f'{idx}. {label}{f"：{url}" if url else ""}', styles['small']))
    return story

@router.get('/{report_id}/export')
async def export_report(report_id: int, format: str = 'md', db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    report = await db.get(Report, report_id)
    if not report or report.user_id != user.id:
        raise HTTPException(404, '报告不存在或无权访问，请刷新报告列表后重试。')
    content = loads(report.content, {})
    download_title = report.title or f'report-{report.id}'
    markdown = _build_export_markdown(content, download_title)
    if format == 'docx':
        doc = Document()
        _add_report_docx(doc, content, download_title)
        bio = BytesIO(); doc.save(bio)
        return Response(bio.getvalue(), media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers=_download_headers(download_title, 'docx'))
    if format == 'pdf':
        font_name = 'Helvetica'
        for font_path in [
            '/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf',
            '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
            '/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc',
        ]:
            try:
                pdfmetrics.registerFont(TTFont('NotoCJK', font_path))
                font_name = 'NotoCJK'
                break
            except Exception:
                pass
        bio = BytesIO()
        pdf_doc = SimpleDocTemplate(
            bio,
            pagesize=A4,
            leftMargin=16 * mm,
            rightMargin=16 * mm,
            topMargin=18 * mm,
            bottomMargin=18 * mm,
            title=download_title,
        )
        pdf_doc.build(_build_pdf_story(content, download_title, font_name))
        return Response(bio.getvalue(), media_type='application/pdf', headers=_download_headers(download_title, 'pdf'))
    return Response(('\ufeff' + markdown).encode('utf-8'), media_type='text/markdown; charset=utf-8', headers=_download_headers(download_title, 'md'))
